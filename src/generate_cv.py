import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF

# Load data from JSON file
with open('data.json', 'r') as file:
    cv_data = json.load(file)

# Create DOCX file
def create_docx(cv_data):
    doc = Document()

    # Title
    title = doc.add_heading(level=0)
    run = title.add_run(cv_data["name"])
    run.bold = True
    run.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact Information
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_info.add_run(f'Phone: {cv_data["phone"]} | ')
    contact_info.add_run(f'Email: {cv_data["email"]} | ')
    linkedin_run = contact_info.add_run("LinkedIn: ")
    linkedin_run.font.color.rgb = RGBColor(0, 0, 255)
    linkedin_run.font.underline = True
    contact_info.add_run(cv_data["linkedin"])

    # About Me
    doc.add_heading("About Me", level=1)
    doc.add_paragraph(cv_data["aboutMe"])

    # Objective
    doc.add_heading("Objective", level=1)
    doc.add_paragraph(cv_data["objective"])

    # Work Experience
    doc.add_heading("Work Experience", level=1)
    for job in cv_data["workExperience"]:
        job_paragraph = doc.add_paragraph()
        job_paragraph.add_run(f'{job["title"]}\n').bold = True
        job_paragraph.add_run(f'{job["date"]}\n').italic = True
        for responsibility in job["responsibilities"]:
            job_paragraph.add_run(f'- {responsibility}\n')

    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph(cv_data["education"])

    # Programming Languages
    doc.add_heading("Programming Languages", level=1)
    languages_paragraph = doc.add_paragraph()
    languages_paragraph.add_run(", ".join(cv_data["programmingLanguages"]))

    # Additional Skills
    doc.add_heading("Additional Skills", level=1)
    skills_paragraph = doc.add_paragraph()
    skills_paragraph.add_run(", ".join(cv_data["additionalSkills"]))

    # Languages
    doc.add_heading("Languages", level=1)
    for lang, level in cv_data["languages"].items():
        doc.add_paragraph(f'{lang}: {level}')

    # References
    doc.add_paragraph()
    doc.add_paragraph("References available on request")

    # Save the document
    doc.save("Wafi_Hassan_CV.docx")

# Create PDF file
class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, "Curriculum Vitae", 0, 1, "C")

    def chapter_title(self, title):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, title, 0, 1, "L")
        self.ln(4)

    def chapter_body(self, body):
        self.set_font("Arial", "", 12)
        self.multi_cell(0, 10, body)
        self.ln()

def create_pdf(cv_data):
    pdf = PDF()
    pdf.add_page()

    # Add content to the PDF
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, cv_data["name"], 0, 1, "C")
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 10, f'Phone: {cv_data["phone"]}', 0, 1, "C")
    pdf.cell(0, 10, f'Email: {cv_data["email"]}', 0, 1, "C")
    pdf.cell(0, 10, f'LinkedIn: {cv_data["linkedin"]}', 0, 1, "C")

    pdf.ln(10)
    pdf.chapter_title("About Me")
    pdf.chapter_body(cv_data["aboutMe"])

    pdf.chapter_title("Objective")
    pdf.chapter_body(cv_data["objective"])

    pdf.chapter_title("Work Experience")
    for job in cv_data["workExperience"]:
        job_title = f'{job["title"]} ({job["date"]})'
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 10, job_title, 0, 1)
        pdf.set_font("Arial", "", 12)
        for responsibility in job["responsibilities"]:
            pdf.multi_cell(0, 10, f'- {responsibility}')
        pdf.ln(5)

    pdf.chapter_title("Education")
    pdf.chapter_body(cv_data["education"])

    pdf.chapter_title("Programming Languages")
    skills = ', '.join(cv_data["programmingLanguages"])
    pdf.chapter_body(skills)

    pdf.chapter_title("Additional Skills")
    additional_skills = ', '.join(cv_data["additionalSkills"])
    pdf.chapter_body(additional_skills)

    pdf.chapter_title("Languages")
    languages = '\n'.join([f'{lang}: {level}' for lang, level in cv_data["languages"].items()])
    pdf.chapter_body(languages)

    # Save the PDF
    pdf.output("Wafi_Hassan_CV.pdf")

# Generate DOCX and PDF files
create_docx(cv_data)
create_pdf(cv_data)
